from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "panduan-pengguna" / "Altora-Panduan-Pengguna.docx"
SCREENSHOTS = ROOT / "docs" / "panduan-pengguna" / "screenshots"

NAVY = "082145"
PURPLE = "A66DD4"
MAGENTA = "C05BC8"
TEXT = "102343"
SECONDARY = "647087"
MUTED = "9099AA"
BORDER = "DDD9E5"
TINT = "F8F2FC"
INFO = "EAF2FC"
WARNING = "FFF4DE"
SUCCESS = "E8F6EE"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def border_bottom(paragraph, color=BORDER):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    p_pr.append(borders)


def set_cell_padding(cell, top=110, start=150, bottom=110, end=150):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_col_width(cell, width_cm):
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is not None:
        tc_w.set(qn("w:w"), str(int(width_cm * 567)))
        tc_w.set(qn("w:type"), "dxa")


def add_run(paragraph, text, bold=False, color=TEXT, size=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Aptos"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)
    return run


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(4)
    add_run(p, text, color=TEXT, size=10.5)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(5)
    add_run(p, text, color=TEXT, size=10.5)
    return p


def add_callout(doc, title, body, color=INFO):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_padding(cell, 160, 200, 160, 200)
    shade(cell, color)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    add_run(p, title, bold=True, color=NAVY, size=10.5)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    add_run(p2, body, color=TEXT, size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_table_header(header)
    for i, name in enumerate(headers):
        cell = header.cells[i]
        set_col_width(cell, widths[i])
        shade(cell, NAVY)
        set_cell_padding(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run(p, name, bold=True, color="FFFFFF", size=9.5)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            set_col_width(cell, widths[i])
            set_cell_padding(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_run(p, value, color=TEXT, size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_screenshot(doc, file_name, title, width):
    add_heading(doc, title, 2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(SCREENSHOTS / file_name), width=Inches(width))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(12)
    add_run(caption, "Contoh tampilan dengan data demo Kopi Nusantara.", color=SECONDARY, size=9)


def add_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p, "ALTORA  |  Panduan Pengguna", bold=True, color=NAVY, size=8.5)
    border_bottom(p, "D9DEE7")

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Altora - Panduan Pengguna Operasional", color=SECONDARY, size=8.5)
    p.add_run("  |  ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    p._p.append(field)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.72)
section.bottom_margin = Inches(0.65)
section.left_margin = Inches(0.75)
section.right_margin = Inches(0.75)
add_header_footer(section)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Aptos"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(TEXT)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.12
for level, size, color in ((1, 18, NAVY), (2, 13.5, NAVY), (3, 11.5, PURPLE)):
    style = styles[f"Heading {level}"]
    style.font.name = "Aptos Display"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos Display")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.page_break_before = False

# Cover
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(78)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "ALTORA", bold=True, color=NAVY, size=28)
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p2, "Panduan Pengguna Operasional", bold=True, color=PURPLE, size=24)
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p3, "Kasir, stok, pelanggan, tim, dan laporan untuk usaha Indonesia.", color=SECONDARY, size=12)
doc.add_paragraph()
add_callout(doc, "Untuk siapa panduan ini?", "Owner, manager, kasir, staff gudang, dan staff layanan pada Altora Cafe, Toko/Retail, Laundry, Counter, Barbershop, serta Company.", TINT)
doc.add_paragraph()
meta = doc.add_table(rows=3, cols=2)
meta.autofit = False
for row, label, value in zip(meta.rows, ["Versi", "Bahasa", "Tujuan"], ["Panduan operasional", "Indonesia", "Siap dibagikan ke tim usaha"]):
    set_col_width(row.cells[0], 4.2)
    set_col_width(row.cells[1], 11.0)
    shade(row.cells[0], NAVY)
    set_cell_padding(row.cells[0]); set_cell_padding(row.cells[1])
    add_run(row.cells[0].paragraphs[0], label, bold=True, color="FFFFFF", size=10)
    add_run(row.cells[1].paragraphs[0], value, color=TEXT, size=10)
doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# Start
add_heading(doc, "Mulai dengan benar", 1)
add_callout(doc, "Prinsip kerja Altora", "Kasir mencatat transaksi, owner mengatur dan menyetujui, manager menjaga operasional. Setiap orang memakai akun sendiri agar shift, stok, absensi, dan audit dapat ditelusuri.")
add_heading(doc, "Setup pertama oleh Owner", 2)
for item in [
    "Pengaturan > Bisnis: isi nama usaha, pajak, footer struk, aturan poin, dan QRIS statis.",
    "Pengaturan > Outlet: buat seluruh cabang, alamat, dan lebar kertas struk 58 atau 80 mm.",
    "Pengaturan > Karyawan: buat akun Owner, Manager, dan Staff; hubungkan ke outlet kerja mereka.",
    "Produk: buat kategori, produk/jasa, harga jual, modal, SKU/barcode, stok awal, dan satuan.",
    "Untuk karton ke biji, buat satuan dan konversi. Aktifkan serial untuk elektronik; batch/expired untuk produk berumur simpan.",
    "Lakukan satu transaksi dan satu tutup shift percobaan sebelum hari operasional pertama.",
]: add_numbered(doc, item)
add_heading(doc, "Hak akses", 2)
add_table(doc, ["Peran", "Tanggung jawab utama"], [
    ("Owner", "Bisnis, outlet, staf, modul, approval, laporan, dan akses penuh."),
    ("Manager", "Operasional cabang, stok, jadwal, pesanan, dan approval yang diizinkan."),
    ("Staff", "Kasir, pelayanan, absensi, dan tugas yang diberikan."),
], [3.0, 12.2])

add_heading(doc, "Peta menu cepat", 2)
add_table(doc, ["Kebutuhan", "Menu"], [
    ("Menjual produk atau jasa", "Kasir"),
    ("Riwayat, void, retur, struk", "Kasir > Riwayat"),
    ("Produk, barcode, stok, serial", "Produk atau Inventory"),
    ("Supplier, PO, barang datang", "Supplier, Purchase Order, Barang Masuk"),
    ("Pindah atau hitung stok", "Transfer Stok, Stock Count"),
    ("Member, poin, voucher", "Member, Voucher"),
    ("Absensi, jadwal, target", "Tim atau Absensi"),
    ("Uang, biaya, laporan", "Uang, Laporan, atau Finance"),
], [6.2, 9.0])

# POS
add_heading(doc, "Kasir dan transaksi", 1)
add_screenshot(doc, "kasir-desktop.png", "Tampilan Kasir - Desktop", 6.85)
add_heading(doc, "SOP buka shift", 2)
for item in [
    "Login memakai akun sendiri, lalu buka Kasir.",
    "Tekan Buka Shift dan masukkan modal awal yang benar-benar ada di laci.",
    "Periksa printer, scanner, produk aktif, stok, dan koneksi sebelum menerima pelanggan.",
]: add_numbered(doc, item)
add_callout(doc, "Penting", "Uang di laci hanya kas tunai. QRIS, transfer, e-wallet, dan channel delivery adalah penerimaan digital. Jangan menggabungkannya saat tutup shift.", WARNING)
add_heading(doc, "SOP membuat transaksi", 2)
for item in [
    "Scan barcode atau QR produk. Jika scanner gagal, cari berdasarkan SKU atau nama produk.",
    "Pilih varian, topping, serial/IMEI, dan jumlah bila diperlukan.",
    "Pilih jenis pesanan: Dine-in, Takeaway, atau Delivery. Untuk Delivery pilih channel yang tepat.",
    "Pilih member sebelum bayar agar poin, deposit, stamp, atau voucher tercatat.",
    "Pilih metode bayar: Tunai, QRIS, Transfer, E-Wallet, Saldo Member, Voucher, atau split payment.",
    "Tekan Bayar satu kali. Bila aplikasi lambat, tunggu respons dan jangan menekan berulang.",
]: add_numbered(doc, item)
add_heading(doc, "SOP tutup shift", 2)
for item in [
    "Selesaikan transaksi dan sinkronkan antrean offline bila ada.",
    "Buka Kasir > Tutup Shift, lalu hitung uang fisik di laci.",
    "Bandingkan uang fisik dengan expected cash. Periksa kas awal, retur tunai, cashback, dan gesek tunai bila ada selisih.",
    "Isi catatan alasan selisih. Jangan mengubah angka untuk memaksa hasil cocok.",
]: add_numbered(doc, item)

# Mobile tablet
add_heading(doc, "Kasir di tablet dan ponsel", 1)
add_screenshot(doc, "kasir-tablet.png", "Tampilan Kasir - Tablet", 5.25)
add_screenshot(doc, "kasir-mobile.png", "Tampilan Kasir - Mobile", 2.65)
add_callout(doc, "Cara memakai mobile", "Gunakan pencarian atau scanner di atas, tambah produk dengan tombol plus, lalu buka invoice dari bagian bawah layar. Pastikan area kamera dan tombol utama tidak tertutup notch atau navigasi perangkat.", INFO)

# inventory
add_heading(doc, "Inventory dan pembelian", 1)
add_callout(doc, "Aturan emas", "Barang datang tidak boleh hanya dicatat di chat. Jumlah, harga modal, QC, satuan, batch, dan tanggal kedaluwarsa harus masuk ke Altora.", WARNING)
add_table(doc, ["Kejadian", "Prosedur di Altora"], [
    ("Pesan barang", "Buat Purchase Order ke supplier. PO belum menambah stok."),
    ("Barang tiba", "Catat di Barang Masuk, lakukan QC, masukkan jumlah, harga modal, batch/expired jika perlu."),
    ("Pindah antar cabang", "Gunakan Transfer Stok. Cabang tujuan wajib menerima transfer."),
    ("Stok fisik berbeda", "Gunakan Stock Count, verifikasi selisih, lalu simpan alasan."),
    ("Rusak, hilang, expired", "Catat di Waste atau penyesuaian stok; jangan menghapus riwayat."),
], [4.4, 10.8])
add_heading(doc, "Barang dengan satuan, expired, dan serial", 2)
for item in [
    "Gunakan satuan dasar untuk barang yang dijual, misalnya biji, botol, gram, atau liter.",
    "Buat konversi bila membeli karton/dus tetapi menjual per biji. Contoh: 1 dus = 24 botol.",
    "Untuk makanan, minuman, obat, dan kosmetik: input batch, tanggal produksi, dan tanggal kedaluwarsa saat barang diterima.",
    "Gunakan FEFO untuk produk berexpired: jual batch dengan tanggal kedaluwarsa paling dekat terlebih dahulu.",
    "Untuk HP, perangkat elektronik, atau barang bergaransi: aktifkan serial/IMEI dan scan serial saat penerimaan maupun penjualan.",
]: add_numbered(doc, item)
add_heading(doc, "Alur penerimaan barang", 2)
for item in [
    "Pilih supplier atau PO yang sesuai.",
    "Scan/cari produk lalu masukkan jumlah dan satuan penerimaan.",
    "Input harga modal, batch/expired/serial jika diperlukan.",
    "Lakukan QC: tandai barang rusak atau kurang sebelum penerimaan diselesaikan.",
    "Simpan penerimaan. Stok bertambah dan data menjadi dasar laporan stok.",
]: add_numbered(doc, item)

# Cafe
add_heading(doc, "Altora Cafe", 1)
add_heading(doc, "Setup cafe", 2)
for item in [
    "Buat kategori Kopi, Non-Kopi, Makanan, dan Add-on.",
    "Buat menu, harga modal, varian ukuran, topping, serta resep bahan baku bila inventory dipakai.",
    "Buat meja dan lantai pada Pengaturan > Meja, lalu cetak QR untuk setiap meja.",
    "Aktifkan modul Pemesanan Digital, Member, Promo, dan HR sesuai kebutuhan usaha.",
]: add_numbered(doc, item)
add_heading(doc, "QR meja dan dapur", 2)
for item in [
    "Pelanggan scan QR meja, memilih menu, dan mengirim pesanan.",
    "Pesanan masuk ke Pesanan Meja atau layar dapur. Staff menerima, memproses, lalu menandai siap.",
    "Saat pembayaran, selesaikan lewat pembayaran meja atau Kasir.",
    "Jika batal, batalkan dari pesanan agar reservasi stok dapat dikembalikan sesuai kebijakan usaha.",
]: add_numbered(doc, item)
add_heading(doc, "Delivery, catering, dan multi-cabang", 2)
add_bullet(doc, "Order counter: pilih Takeaway.")
add_bullet(doc, "Gojek, Grab, Shopee Food, kurir toko, atau channel lain: pilih Delivery dan channel yang tepat.")
add_bullet(doc, "Pesanan event: gunakan Pesanan Katering agar DP, transport, biaya barista, dan pelunasan terlacak.")
add_bullet(doc, "Cabang A kurang bahan: buat Transfer Stok dari cabang B, kemudian cabang A menerima transfer. Jangan mengubah stok manual.")

# retail/laundry/services
add_heading(doc, "Altora Toko, Laundry, dan layanan", 1)
add_heading(doc, "Altora Toko / Retail", 2)
for item in [
    "Setiap barang dengan SKU, ukuran, atau harga berbeda harus dibuat sebagai produk/varian yang jelas.",
    "Isi barcode dan cetak label bila barang belum memiliki barcode dari supplier.",
    "Aktifkan expiry untuk makanan/minuman/kosmetik; serial untuk elektronik atau perangkat bergaransi.",
    "Atur stok minimum agar item yang cepat habis muncul sebelum rak kosong.",
    "Jika transaksi sudah selesai, gunakan void atau retur dari riwayat. Jangan menghapus data transaksi.",
]: add_numbered(doc, item)
add_heading(doc, "Altora Laundry", 2)
for item in [
    "Buat layanan Kiloan, Express, Satuan, Dry Clean, dan Setrika pada Pengaturan > Laundry.",
    "Saat order: pilih pelanggan, layanan, berat/jumlah, harga, biaya tambahan, pickup/delivery, estimasi selesai, catatan kondisi, dan DP.",
    "Gunakan status: Diterima > Dicuci > Dikeringkan > Disetrika > Siap > Diambil.",
    "Catat pembayaran bertahap bila pelanggan memberi DP dan melunasi saat barang diambil.",
    "Saat serah terima, cocokkan nomor order, nomor HP, total, dan status pembayaran.",
]: add_numbered(doc, item)
add_heading(doc, "Counter dan Barbershop", 2)
add_bullet(doc, "Counter: bedakan barang stok dan jasa. Jasa tidak mengurangi stok. Gunakan barcode untuk aksesoris dan serial untuk perangkat unik.")
add_bullet(doc, "Barbershop: aktifkan Booking dan Tim. Buat jasa/paket, booking berdasarkan tanggal-jam-staff, lalu selesaikan pembayaran di Kasir.")

# team finance etc
add_heading(doc, "Pelanggan, tim, uang, dan laporan", 1)
add_heading(doc, "Member dan promo", 2)
add_bullet(doc, "Pilih member sebelum pembayaran untuk poin, deposit, voucher, atau stamp.")
add_bullet(doc, "Gunakan promo terjadwal untuk happy hour atau kategori tertentu; diskon manual sebaiknya diberi alasan.")
add_heading(doc, "Tim dan dokumen", 2)
add_bullet(doc, "Manager membuat jadwal per outlet/shift. Staff clock-in dan clock-out dengan foto/lokasi bila modul diaktifkan.")
add_bullet(doc, "Gunakan Dokumen dan e-sign untuk SP, izin, kontrak, dan approval. Atur akses berdasarkan user atau peran.")
add_heading(doc, "Keuangan", 2)
add_bullet(doc, "Catat biaya operasional pada hari kejadian: bahan baku, sewa, listrik, gaji, transport, dan lain-lain.")
add_bullet(doc, "Mode Simple untuk kas masuk/keluar dan insight ringkas.")
add_bullet(doc, "Mode Advanced untuk finance: jurnal, hutang supplier, laba-rugi, neraca saldo, dan tutup buku.")
add_heading(doc, "Checklist akhir hari Owner", 2)
for item in [
    "Pastikan semua kasir sudah tutup shift dan alasan selisih sudah dicatat.",
    "Bandingkan kas fisik dengan laporan kas. Periksa QRIS/digital sebagai kelompok terpisah.",
    "Periksa stok menipis, barang expired/akan expired, transfer yang belum diterima, dan PO yang belum selesai.",
    "Tinjau transaksi void, retur, diskon manual, dan audit log bila ada hal yang tidak biasa.",
    "Pastikan antrean offline sudah tersinkronkan sebelum data browser/perangkat dibersihkan.",
]: add_numbered(doc, item)

# Troubleshooting
add_heading(doc, "Offline, masalah umum, dan keamanan", 1)
add_heading(doc, "Internet putus", 2)
add_bullet(doc, "Jangan menekan Bayar berulang. Transaksi POS dapat masuk antrean perangkat dan akan disinkronkan saat koneksi kembali.")
add_bullet(doc, "Jangan hapus data browser atau berpindah perangkat sebelum antrean hijau/sinkron selesai.")
add_bullet(doc, "Jika gagal karena stok berubah, selesaikan konflik secara manual dengan supervisor.")
add_heading(doc, "QRIS atau barcode bermasalah", 2)
add_table(doc, ["Masalah", "Tindakan"], [
    ("QRIS", "Periksa QR statis di Pengaturan. QR nominal bukan verifikasi mutasi otomatis; kasir tetap mengikuti SOP usaha untuk memastikan pembayaran masuk."),
    ("Barcode", "Periksa izin kamera/scanner, bersihkan lensa, cari SKU/nama, lalu cocokkan produk dengan label."),
    ("Printer", "Periksa kertas, koneksi, dan ukuran struk. Coba cetak ulang satu struk percobaan."),
    ("Stok tidak sesuai", "Gunakan Stock Count dan beri alasan selisih. Jangan mengubah stok secara diam-diam."),
], [3.2, 12.0])
add_heading(doc, "Keamanan akun", 2)
for item in [
    "Jangan membagikan akun Owner atau password kepada staff.",
    "Nonaktifkan akses staff yang sudah keluar atau pindah cabang.",
    "Periksa Pengaturan > Log Audit untuk transaksi, perubahan harga, dan akses mencurigakan.",
    "Ganti secret/API key yang pernah muncul di chat, dokumen, atau screenshot.",
]: add_numbered(doc, item)
add_callout(doc, "Butuh bantuan?", "Catat nomor transaksi/order, outlet, waktu kejadian, dan screenshot pesan error sebelum menghubungi administrator Altora. Informasi itu mempercepat pengecekan.", SUCCESS)

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
